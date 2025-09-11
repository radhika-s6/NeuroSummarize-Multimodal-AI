import cv2
import numpy as np
from PIL import Image
import easyocr
import pytesseract
from transformers import LayoutLMv3Processor, LayoutLMv3ForTokenClassification
import torch
import fitz  # PyMuPDF for PDF handling
import io
import tempfile
import os
import pandas as pd
import json
import zipfile
from pathlib import Path
from src.utils import preprocessing
from src.utils.preprocessing import clean_text

# Add docx import with fallback
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class MultiModalOCR:
    def __init__(self):
        self.easyocr_reader = easyocr.Reader(['en'], verbose=False)
        self.layoutlm_processor = None
        self.layoutlm_model = None
        
        # Supported file formats
        self.supported_formats = {
            'images': ['.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp'],
            'text': ['.txt', '.md', '.rtf'],
            'documents': ['.docx'],
            'pdfs': ['.pdf'],
            'data': ['.csv', '.xlsx', '.xls', '.json', '.jsonl'],
            'archives': ['.zip']
        }

    def detect_file_type(self, filename):
        """Detect file type from extension"""
        ext = Path(filename).suffix.lower()
        for category, extensions in self.supported_formats.items():
            if ext in extensions:
                return category
        return 'unknown'

    def preprocess_image(self, image_path):
        """Preprocess image for better OCR results"""
        img = cv2.imread(str(image_path))
        if img is None:
            raise FileNotFoundError(f"Image not found or unreadable: {image_path}")
        
        # Convert to grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        # Noise removal
        denoised = cv2.medianBlur(gray, 5)
        # Thresholding
        thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        # Morphological operations
        kernel = np.ones((1, 1), np.uint8)
        processed = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
        return processed

    def extract_text_from_pdf_direct(self, pdf_path):
        """Extract text directly from PDF"""
        try:
            doc = fitz.open(pdf_path)
            text_content = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                if text.strip():
                    text_content.append(f"=== Page {page_num + 1} ===\n{text}")
            
            doc.close()
            
            if text_content:
                return {
                    'text': '\n\n'.join(text_content),
                    'confidence': 1.0,
                    'method': 'direct_pdf_extraction',
                    'pages_processed': len(text_content)
                }
            return None
        except Exception as e:
            return None

    def pdf_to_images(self, pdf_path, output_dir=None):
        """Convert PDF pages to images"""
        if output_dir is None:
            output_dir = tempfile.mkdtemp()
        
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        try:
            doc = fitz.open(pdf_path)
            image_paths = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                pix = page.get_pixmap(matrix=mat)
                image_path = output_dir / f"page_{page_num + 1:03d}.png"
                pix.save(str(image_path))
                image_paths.append(str(image_path))
            
            doc.close()
            return image_paths
        except Exception as e:
            raise Exception(f"Failed to convert PDF to images: {str(e)}")

    def extract_from_pdf_ocr(self, pdf_path):
        """Extract text from PDF using OCR"""
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                image_paths = self.pdf_to_images(pdf_path, temp_dir)
                page_results = []
                total_confidence = 0
                
                for i, image_path in enumerate(image_paths):
                    page_result = self.ensemble_extraction_image(image_path)
                    if page_result['final_text'].strip():
                        page_results.append({
                            'page': i + 1,
                            'text': page_result['final_text'],
                            'confidence': page_result['confidence']
                        })
                        total_confidence += page_result['confidence']
                
                if page_results:
                    combined_text = '\n\n'.join([
                        f"=== Page {result['page']} ===\n{result['text']}" 
                        for result in page_results
                    ])
                    avg_confidence = total_confidence / len(page_results)
                    return {
                        'text': combined_text,
                        'confidence': avg_confidence,
                        'method': 'pdf_ocr_extraction',
                        'pages_processed': len(page_results)
                    }
                else:
                    return {
                        'text': '',
                        'confidence': 0,
                        'method': 'pdf_ocr_extraction',
                        'error': 'No text extracted'
                    }
        except Exception as e:
            return {'text': '', 'confidence': 0, 'method': 'pdf_ocr_extraction', 'error': str(e)}

    def extract_from_pdf(self, pdf_path):
        """Comprehensive PDF extraction"""
        try:
            # Try direct extraction first
            direct_result = self.extract_text_from_pdf_direct(pdf_path)
            
            if direct_result and len(direct_result['text'].strip()) > 50:
                return direct_result
            
            # Fall back to OCR
            return self.extract_from_pdf_ocr(pdf_path)
        except Exception as e:
            return {'text': '', 'confidence': 0, 'method': 'pdf_extraction_failed', 'error': str(e)}

    def extract_from_text_file(self, file_path):
        """Extract from text files"""
        try:
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    return {
                        'text': content,
                        'confidence': 1.0,
                        'method': 'direct_text',
                        'encoding': encoding
                    }
                except UnicodeDecodeError:
                    continue
            return {'text': '', 'confidence': 0, 'method': 'text_extraction', 'error': 'Encoding error'}
        except Exception as e:
            return {'text': '', 'confidence': 0, 'method': 'text_extraction', 'error': str(e)}

    def extract_from_docx(self, file_path):
        """Extract from DOCX files"""
        if not DOCX_AVAILABLE:
            return {'text': '', 'confidence': 0, 'method': 'docx_extraction', 'error': 'python-docx not installed'}
        
        try:
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_text.append(row_text)
            
            full_text = '\n\n'.join(paragraphs)
            if table_text:
                full_text += '\n\nTABLES:\n' + '\n'.join(table_text)
            
            return {
                'text': full_text,
                'confidence': 1.0,
                'method': 'docx_extraction',
                'paragraphs': len(paragraphs),
                'tables': len(table_text)
            }
        except Exception as e:
            return {'text': '', 'confidence': 0, 'method': 'docx_extraction', 'error': str(e)}

    def extract_from_csv(self, file_path):
        """Extract from CSV files"""
        try:
            # Try different separators
            for sep in [',', ';', '\t', '|']:
                try:
                    df = pd.read_csv(file_path, sep=sep, encoding='utf-8')
                    if len(df.columns) > 1:
                        break
                except:
                    continue
            else:
                df = pd.read_csv(file_path)
            
            text_parts = [f"CSV FILE: {Path(file_path).name}"]
            text_parts.append(f"COLUMNS: {', '.join(df.columns)}")
            text_parts.append(f"ROWS: {len(df)}")
            text_parts.append("\nDATA PREVIEW:")
            
            for idx, row in df.head(10).iterrows():
                row_text = ' | '.join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                text_parts.append(f"Row {idx + 1}: {row_text}")
            
            return {
                'text': '\n'.join(text_parts),
                'confidence': 1.0,
                'method': 'csv_parsing',
                'rows': len(df),
                'columns': len(df.columns)
            }
        except Exception as e:
            return {'text': '', 'confidence': 0, 'method': 'csv_parsing', 'error': str(e)}

    def extract_from_excel(self, file_path):
        """Extract from Excel files"""
        try:
            excel_file = pd.ExcelFile(file_path)
            text_parts = [f"EXCEL FILE: {Path(file_path).name}"]
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text_parts.append(f"\nSHEET: {sheet_name}")
                text_parts.append(f"COLUMNS: {', '.join(df.columns)}")
                text_parts.append(f"ROWS: {len(df)}")
                
                for idx, row in df.head(5).iterrows():
                    row_text = ' | '.join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                    text_parts.append(f"  {row_text}")
            
            return {
                'text': '\n'.join(text_parts),
                'confidence': 1.0,
                'method': 'excel_parsing',
                'sheets': len(excel_file.sheet_names)
            }
        except Exception as e:
            return {'text': '', 'confidence': 0, 'method': 'excel_parsing', 'error': str(e)}

    def extract_from_json(self, file_path):
        """Extract from JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                if file_path.endswith('.jsonl'):
                    data = [json.loads(line) for line in f if line.strip()]
                else:
                    data = json.load(f)
            
            text_parts = [f"JSON FILE: {Path(file_path).name}"]
            
            if isinstance(data, list):
                text_parts.append(f"ARRAY LENGTH: {len(data)}")
                text_parts.append("SAMPLE ITEMS:")
                for i, item in enumerate(data[:5]):
                    text_parts.append(f"Item {i + 1}: {json.dumps(item, indent=2)}")
            elif isinstance(data, dict):
                text_parts.append("JSON OBJECT:")
                text_parts.append(json.dumps(data, indent=2))
            else:
                text_parts.append(f"JSON VALUE: {data}")
            
            return {
                'text': '\n'.join(text_parts),
                'confidence': 1.0,
                'method': 'json_parsing',
                'json_data': data
            }
        except Exception as e:
            return {'text': '', 'confidence': 0, 'method': 'json_parsing', 'error': str(e)}

    def extract_with_easyocr(self, image_path):
        """Extract text using EasyOCR"""
        try:
            processed_img = self.preprocess_image(image_path)
            results = self.easyocr_reader.readtext(processed_img)
            text_blocks = []
            confidence_scores = []
            
            for (bbox, text, confidence) in results:
                if confidence > 0.5:
                    text_blocks.append(text)
                    confidence_scores.append(confidence)
            
            return {
                'text': ' '.join(text_blocks),
                'confidence': np.mean(confidence_scores) if confidence_scores else 0,
                'method': 'easyocr'
            }
        except Exception as e:
            return {'text': '', 'confidence': 0, 'method': 'easyocr', 'error': str(e)}

    def extract_with_tesseract(self, image_path):
        """Extract text using Tesseract OCR"""
        try:
            processed_img = self.preprocess_image(image_path)
            text = pytesseract.image_to_string(processed_img, config='--psm 6')
            
            # Get confidence scores
            data = pytesseract.image_to_data(processed_img, output_type=pytesseract.Output.DICT)
            confidences = [int(conf) for conf in data['conf'] if conf != '-1']
            avg_confidence = np.mean(confidences) / 100 if confidences else 0
            
            return {
                'text': text.strip(),
                'confidence': avg_confidence,
                'method': 'tesseract'
            }
        except Exception as e:
            return {'text': '', 'confidence': 0, 'method': 'tesseract', 'error': str(e)}

    def ensemble_extraction_image(self, image_path):
        """Ensemble extraction for images only"""
        results = {
            'easyocr': self.extract_with_easyocr(image_path),
            'tesseract': self.extract_with_tesseract(image_path)
        }
        
        # Choose best result
        best_result = max(results.values(), 
                         key=lambda x: x.get('confidence', 0) * len(x.get('text', '')))
        
        return {
            'final_text': best_result.get('text', ''),
            'best_method': best_result.get('method', 'none'),
            'confidence': best_result.get('confidence', 0),
            'all_results': results
        }

    def ensemble_extraction(self, file_path):
        """Universal file extraction method"""
        file_path = Path(file_path)
        file_type = self.detect_file_type(file_path.name)
        
        try:
            if file_type == 'images':
                return self.ensemble_extraction_image(file_path)
            elif file_type == 'pdfs':
                result = self.extract_from_pdf(file_path)
                return {
                    'final_text': result.get('text', ''),
                    'best_method': result.get('method', 'pdf_extraction'),
                    'confidence': result.get('confidence', 0),
                    'all_results': {'pdf': result}
                }
            elif file_type == 'text':
                result = self.extract_from_text_file(file_path)
                return {
                    'final_text': result.get('text', ''),
                    'best_method': result.get('method', 'text_extraction'),
                    'confidence': result.get('confidence', 0),
                    'all_results': {'text': result}
                }
            elif file_type == 'documents':
                result = self.extract_from_docx(file_path)
                return {
                    'final_text': result.get('text', ''),
                    'best_method': result.get('method', 'docx_extraction'),
                    'confidence': result.get('confidence', 0),
                    'all_results': {'docx': result}
                }
            elif file_type == 'data':
                ext = file_path.suffix.lower()
                if ext == '.csv':
                    result = self.extract_from_csv(file_path)
                elif ext in ['.xlsx', '.xls']:
                    result = self.extract_from_excel(file_path)
                elif ext in ['.json', '.jsonl']:
                    result = self.extract_from_json(file_path)
                else:
                    return {'final_text': '', 'best_method': 'unsupported', 'confidence': 0, 'all_results': {}}
                
                return {
                    'final_text': result.get('text', ''),
                    'best_method': result.get('method', 'data_extraction'),
                    'confidence': result.get('confidence', 0),
                    'all_results': {'data': result}
                }
            else:
                return {
                    'final_text': '',
                    'best_method': 'unsupported_format',
                    'confidence': 0,
                    'all_results': {},
                    'error': f'Unsupported file type: {file_type}'
                }
        except Exception as e:
            return {
                'final_text': '',
                'best_method': 'extraction_failed',
                'confidence': 0,
                'all_results': {},
                'error': str(e)
            }

    def extract_with_layoutlm(self, image_path):
        """Extract text using LayoutLMv3 (placeholder)"""
        return {'text': '', 'confidence': 0, 'method': 'layoutlm', 'error': 'Not implemented'}